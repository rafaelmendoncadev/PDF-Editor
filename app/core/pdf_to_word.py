"""
PDF to Word converter — uses PyMuPDF to extract content and python-docx to build
a well-formatted .docx preserving text, tables, and basic styling.

Much better output than pdf2docx for tabular/structured PDFs.
"""

from __future__ import annotations

import logging
from typing import List, Tuple

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


def convert_pdf_to_word(pdf_path: str, docx_path: str) -> None:
    """
    Convert *pdf_path* to a Word document at *docx_path*.

    Uses PyMuPDF to extract tables and text blocks, then builds
    a .docx with python-docx preserving structure and formatting.
    """
    try:
        pdf = fitz.open(pdf_path)
        doc = Document()

        # Set narrow margins for better use of space
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        for page_idx in range(pdf.page_count):
            page = pdf.load_page(page_idx)

            if page_idx > 0:
                doc.add_page_break()

            _process_page(doc, page)

        doc.save(docx_path)
        pdf.close()
        logger.info("Converted %s → %s", pdf_path, docx_path)

    except ImportError as exc:
        raise RuntimeError(
            "Dependência faltando. Execute:\n"
            "  pip install PyMuPDF python-docx"
        ) from exc
    except Exception as exc:
        raise RuntimeError(f"Falha na conversão: {exc}") from exc


# ----------------------------------------------------------------------
# Internal helpers
# ----------------------------------------------------------------------

def _process_page(doc: Document, page: fitz.Page) -> None:
    """Extract content from a single PDF page and add to the Word doc."""

    # 1) Find tables on this page
    tables = page.find_tables()
    table_rects: List[fitz.Rect] = []
    for tab in tables:
        table_rects.append(fitz.Rect(tab.bbox))

    # 2) Get all text blocks (type 0 = text)
    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

    # Separate text blocks that are NOT inside any table
    text_blocks = []
    for blk in blocks:
        if blk.get("type") != 0:
            continue
        blk_rect = fitz.Rect(blk["bbox"])
        inside_table = any(
            tr.intersects(blk_rect) for tr in table_rects
        )
        if not inside_table:
            text_blocks.append(blk)

    # 3) Build a list of elements sorted by vertical position
    elements: List[Tuple[float, str, object]] = []  # (y, kind, data)

    for blk in text_blocks:
        elements.append((blk["bbox"][1], "text", blk))

    for tab in tables:
        elements.append((tab.bbox[1], "table", tab))

    elements.sort(key=lambda e: e[0])

    # 4) Render each element into the docx
    for _, kind, data in elements:
        if kind == "text":
            _add_text_block(doc, data)
        elif kind == "table":
            _add_table(doc, data)


def _add_text_block(doc: Document, block: dict) -> None:
    """Add a text block (may contain multiple lines/spans) to the doc."""
    for line in block.get("lines", []):
        spans = line.get("spans", [])
        if not spans:
            continue

        # Combine all spans in this line into one paragraph
        full_text = "".join(sp.get("text", "") for sp in spans).strip()
        if not full_text:
            continue

        para = doc.add_paragraph()

        # Detect alignment heuristically from first span x-position
        page_width = 595  # A4 approx
        first_x = spans[0].get("origin", [0])[0] if spans[0].get("origin") else spans[0]["bbox"][0]
        text_width = spans[-1]["bbox"][2] - spans[0]["bbox"][0]
        center_x = first_x + text_width / 2

        if abs(center_x - page_width / 2) < 40:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for sp in spans:
            text = sp.get("text", "")
            if not text:
                continue

            run = para.add_run(text)
            size = sp.get("size", 11)
            run.font.size = Pt(max(6, min(size, 36)))

            # Bold detection
            flags = sp.get("flags", 0)
            if flags & 2 ** 4:  # bit 4 = bold
                run.bold = True
            if flags & 2 ** 1:  # bit 1 = italic
                run.italic = True

            # Color
            color_int = sp.get("color", 0)
            if color_int and color_int != 0:
                r = (color_int >> 16) & 0xFF
                g = (color_int >> 8) & 0xFF
                b = color_int & 0xFF
                run.font.color.rgb = RGBColor(r, g, b)

            # Font name
            font_name = sp.get("font", "")
            if font_name:
                # Map common PDF font names to Word equivalents
                clean = font_name.split(",")[0].split("-")[0].strip()
                if clean:
                    run.font.name = clean


def _add_table(doc: Document, table: fitz.table.Table) -> None:
    """Add a PDF table to the Word document."""
    try:
        data = table.extract()
    except Exception:
        return

    if not data or not data[0]:
        return

    num_rows = len(data)
    num_cols = len(data[0])

    if num_cols == 0:
        return

    # Create Word table
    word_table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(data):
        row_cells = word_table.rows[r_idx].cells
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= len(row_cells):
                break
            text = (cell_text or "").strip()
            cell = row_cells[c_idx]
            cell.text = ""

            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run(text)
            run.font.size = Pt(10)

            # Bold first row (header)
            if r_idx == 0:
                run.bold = True

    # Add a small space after the table
    doc.add_paragraph("")
